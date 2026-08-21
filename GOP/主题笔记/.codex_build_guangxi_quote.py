from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"D:\mygithub\mynote\DescriptiveStatistic\GOP\主题笔记")
SOURCE = ROOT / "广西报价.md"
OUTPUT = ROOT / "广西报价-正式版.docx"

# Design system:
# - preset: standard_business_brief
# - named override: cn_quotation_a4
#   A4 portrait; 1.8 cm margins; Microsoft YaHei for Chinese readability.
PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_CM = 1.8
CONTENT_WIDTH_CM = PAGE_WIDTH_CM - 2 * MARGIN_CM
CONTENT_WIDTH_DXA = round(CONTENT_WIDTH_CM / 2.54 * 1440)
TABLE_INDENT_DXA = 120

FONT_LATIN = "Arial"
FONT_EAST_ASIA = "Microsoft YaHei"

NAVY = "17365D"
BLUE = "2E75B6"
PALE_BLUE = "EAF2F8"
LIGHT_BLUE = "D9EAF7"
PALE_GRAY = "F2F4F7"
MID_GRAY = "D8DEE6"
TEXT_GRAY = "5B6573"
DARK = "1F2328"
WHITE = "FFFFFF"
GOLD = "C28A19"
PALE_GOLD = "FFF7E3"
GREEN = "2E7D57"
PALE_GREEN = "E9F5EF"
RED = "9B1C1C"


@dataclass
class Module:
    number: int
    name: str
    reference: str
    days_low: int
    days_high: int
    fee_low: float
    fee_high: float
    requirement: str
    credential: str


def clean_markdown_text(text: str) -> str:
    text = text.replace("\u00a0", " ").replace("\u200b", "")
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def parse_range(text: str, cast=float):
    values = re.findall(r"\d+(?:\.\d+)?", text)
    if len(values) < 2:
        raise ValueError(f"Cannot parse range from: {text!r}")
    return cast(values[0]), cast(values[1])


def parse_modules(source: str) -> list[Module]:
    headings = list(re.finditer(r"^### 模块(\d+)：(.*)$", source, re.M))
    modules: list[Module] = []
    section_end = source.find("\n## 三、", headings[-1].end())
    if section_end < 0:
        section_end = len(source)

    for idx, match in enumerate(headings):
        end = headings[idx + 1].start() if idx + 1 < len(headings) else section_end
        segment = source[match.end():end]
        heading = clean_markdown_text(match.group(2))

        relation_match = re.match(r"^(.*?)（对应(.*?)）$", heading)
        if relation_match:
            name = relation_match.group(1).strip()
            reference = relation_match.group(2).strip()
        else:
            name = heading
            reference = "—"

        req_match = re.search(
            r"\*\*技术要求原文：\*\*\s*(.*?)\s*\*\*估算人天：\*\*",
            segment,
            re.S,
        )
        days_match = re.search(r"\*\*估算人天：\*\*\s*([^\r\n]+)", segment)
        fee_match = re.search(r"\*\*估算费用：\*\*\s*([^\r\n]+)", segment)
        if not (req_match and days_match and fee_match):
            raise ValueError(f"Module {match.group(1)} could not be parsed")

        req_lines = []
        for line in req_match.group(1).splitlines():
            line = re.sub(r"^\s*>\s?", "", line).strip()
            if line:
                req_lines.append(clean_markdown_text(line))
        requirement = "\n".join(req_lines)

        credential = ""
        cred_match = re.search(r"[\(（]上述所有内容需提供(.*?)[\)）]\s*$", requirement, re.S)
        if cred_match:
            credential = "上述所有内容需提供" + cred_match.group(1).strip()
            requirement = requirement[:cred_match.start()].rstrip("；。 \n")

        days_low, days_high = parse_range(days_match.group(1), int)
        fee_low, fee_high = parse_range(fee_match.group(1), float)

        modules.append(
            Module(
                number=int(match.group(1)),
                name=name,
                reference=reference,
                days_low=days_low,
                days_high=days_high,
                fee_low=fee_low,
                fee_high=fee_high,
                requirement=requirement,
                credential=credential,
            )
        )

    if len(modules) != 27:
        raise ValueError(f"Expected 27 modules, got {len(modules)}")
    return modules


def set_run_font(
    run,
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
    east_asia: str = FONT_EAST_ASIA,
    latin: str = FONT_LATIN,
):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:cs"), latin)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, size: float, bold=False, color=DARK):
    style.font.name = FONT_LATIN
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name, attrs in edges.items():
        edge = tc_borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            tc_borders.append(edge)
        for key, value in attrs.items():
            edge.set(qn(f"w:{key}"), str(value))


def set_paragraph_shading(paragraph, fill: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_paragraph_left_border(paragraph, color=BLUE, size=18, space=8):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = p_bdr.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        p_bdr.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=TABLE_INDENT_DXA):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        prevent_row_split(row)
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Cm(width / 1440 * 2.54)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color=MID_GRAY, size=5):
    for row in table.rows:
        for cell in row.cells:
            attrs = {"val": "single", "sz": size, "color": color, "space": "0"}
            set_cell_border(
                cell,
                top=attrs,
                bottom=attrs,
                start=attrs,
                end=attrs,
                insideH=attrs,
                insideV=attrs,
            )


def set_cell_text(
    cell,
    text: str,
    *,
    size=9.3,
    bold=False,
    color=DARK,
    align=WD_ALIGN_PARAGRAPH.LEFT,
):
    p = cell.paragraphs[0]
    p.alignment = align
    set_paragraph_spacing(p, before=0, after=0, line=1.12)
    p.paragraph_format.keep_together = True
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text: str, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, size=16, bold=True, color=NAVY)
    elif level == 2:
        set_run_font(r, size=12.5, bold=True, color=BLUE)
    else:
        set_run_font(r, size=11, bold=True, color=NAVY)
    return p


def add_body(doc, text: str, *, bold_prefix: str | None = None, italic=False, color=DARK):
    p = doc.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=10.5, bold=True, color=color)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=10.5, italic=italic, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, size=10.5, italic=italic, color=color)
    return p


def add_callout(doc, label: str, text: str, fill=PALE_BLUE, accent=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.28)
    p.paragraph_format.right_indent = Cm(0.12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    set_paragraph_shading(p, fill)
    set_paragraph_left_border(p, color=accent, size=20, space=8)
    r1 = p.add_run(label + " ")
    set_run_font(r1, size=10.2, bold=True, color=accent)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.2, color=DARK)
    return p


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=8.5, color=TEXT_GRAY)


def add_numbering_definition(doc, lvl_text="%1.", start=1, left=540, hanging=270):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(x.get(qn("w:abstractNumId")))
        for x in numbering.findall(qn("w:abstractNum"))
        if x.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(x.get(qn("w:numId")))
        for x in numbering.findall(qn("w:num"))
        if x.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids or [0]) + 1
    num_id = max(num_ids or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start_node = OxmlElement("w:start")
    start_node.set(qn("w:val"), str(start))
    lvl.append(start_node)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl.append(num_fmt)
    lvl_text_node = OxmlElement("w:lvlText")
    lvl_text_node.set(qn("w:val"), lvl_text)
    lvl.append(lvl_text_node)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "space")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left))
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int, level=0):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)


def add_numbered_item(doc, text: str, num_id: int, *, size=10.2, bold=False):
    p = doc.add_paragraph(style="Normal")
    apply_numbering(p, num_id)
    p.paragraph_format.left_indent = None
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.18
    run = p.add_run(text.strip())
    set_run_font(run, size=size, bold=bold, color=DARK)
    return p


def add_requirement_paragraph(doc, text: str, *, bold=False):
    p = doc.add_paragraph(style="Requirement Text")
    p.paragraph_format.keep_together = False
    r = p.add_run(text.strip())
    set_run_font(r, size=9.6, bold=bold, color=DARK)
    return p


def split_numbered(text: str, pattern: str):
    matches = list(re.finditer(pattern, text))
    if not matches:
        return text.strip(), []
    intro = text[:matches[0].start()].strip()
    items = []
    for i, m in enumerate(matches):
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        number = int(m.group(1))
        item = text[m.end():end].strip(" ；;")
        items.append((number, item))
    return intro, items


def add_number_sequence(doc, items: list[tuple[int, str]], lvl_text="（%1）"):
    current_num_id = None
    previous = None
    for number, text in items:
        if current_num_id is None or previous is None or number != previous + 1:
            current_num_id = add_numbering_definition(
                doc,
                lvl_text=lvl_text,
                start=number,
                left=520,
                hanging=300,
            )
        add_numbered_item(doc, text, current_num_id, size=9.6)
        previous = number


def add_complex_requirement(doc, text: str):
    lines = [x.strip() for x in text.splitlines() if x.strip()]
    if len(lines) > 1 and all(re.match(r"^▲?\d+\s+", x) for x in lines):
        items = []
        for line in lines:
            match = re.match(r"^(▲?)(\d+)\s+(.*)$", line)
            marker = "▲ " if match.group(1) else ""
            items.append((int(match.group(2)), marker + match.group(3)))
        add_number_sequence(doc, items, lvl_text="%1.")
        return

    top_matches = list(re.finditer(r"(?<![，；、])([一二三四五六七八九十]+)、", text))
    if top_matches:
        intro = text[:top_matches[0].start()].strip()
        if intro:
            add_requirement_paragraph(doc, intro)
        for idx, match in enumerate(top_matches):
            end = top_matches[idx + 1].start() if idx + 1 < len(top_matches) else len(text)
            segment = text[match.end():end].strip()
            item_match = re.search(r"(?<!\d)(\d+)[.、]", segment)
            if item_match:
                title = segment[:item_match.start()].strip(" ：:；;")
                body = segment[item_match.start():]
            else:
                title = segment.strip(" ：:；;")
                body = ""
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(f"{match.group(1)}、{title}")
            set_run_font(r, size=10, bold=True, color=NAVY)
            if body:
                _, items = split_numbered(body, r"(?<!\d)(\d+)[.、]")
                add_number_sequence(doc, items, lvl_text="%1.")
        return

    intro, items = split_numbered(text, r"（(\d+)）")
    if items:
        if intro:
            add_requirement_paragraph(doc, intro)
        add_number_sequence(doc, items, lvl_text="（%1）")
        return

    if len(lines) > 1:
        for line in lines:
            add_requirement_paragraph(doc, line)
    else:
        add_requirement_paragraph(doc, text)


def add_docx_core_properties(doc):
    props = doc.core_properties
    props.title = "广西项目软件开发报价方案"
    props.subject = "车辆与装备综合管理系统软件定制开发及配套服务报价"
    props.author = "报价编制组"
    props.keywords = "广西, 报价, 软件开发, 车辆管理, 系统集成"
    props.comments = "依据广西报价.md整理；金额按分项明细复核。"


def configure_styles(doc):
    normal = doc.styles["Normal"]
    set_style_font(normal, 10.5, color=DARK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = doc.styles["Title"]
    set_style_font(title, 27, bold=True, color=NAVY)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)

    subtitle = doc.styles["Subtitle"]
    set_style_font(subtitle, 14, color=TEXT_GRAY)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(10)

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, 16, bold=True, color=NAVY)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, 12.5, bold=True, color=BLUE)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(5)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, 11, bold=True, color=NAVY)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    if "Requirement Text" not in [s.name for s in doc.styles]:
        req = doc.styles.add_style("Requirement Text", 1)
    else:
        req = doc.styles["Requirement Text"]
    set_style_font(req, 9.6, color=DARK)
    req.paragraph_format.left_indent = Cm(0.25)
    req.paragraph_format.right_indent = Cm(0.05)
    req.paragraph_format.space_before = Pt(0)
    req.paragraph_format.space_after = Pt(4)
    req.paragraph_format.line_spacing = 1.18
    req.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def configure_section(section):
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_CM)
    section.bottom_margin = Cm(MARGIN_CM)
    section.left_margin = Cm(MARGIN_CM)
    section.right_margin = Cm(MARGIN_CM)
    section.header_distance = Cm(0.9)
    section.footer_distance = Cm(0.9)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, after=0, line=1.0)
    r = p.add_run("广西项目软件开发报价方案")
    set_run_font(r, size=8.5, bold=True, color=TEXT_GRAY)
    r2 = p.add_run("   |   V1.0")
    set_run_font(r2, size=8.5, color=TEXT_GRAY)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(fp, before=0, after=0, line=1.0)
    r = fp.add_run("第 ")
    set_run_font(r, size=8.5, color=TEXT_GRAY)
    add_page_field(fp)
    r = fp.add_run(" 页")
    set_run_font(r, size=8.5, color=TEXT_GRAY)


def add_cover(doc):
    for _ in range(3):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=12)
    r = p.add_run("QUOTATION")
    set_run_font(r, size=10, bold=True, color=GOLD)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=10)
    r = p.add_run("广西项目软件开发报价方案")
    set_run_font(r, size=28, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=30)
    r = p.add_run("车辆与装备综合管理系统 · 软件定制开发及配套服务")
    set_run_font(r, size=14, color=TEXT_GRAY)

    meta = doc.add_table(rows=4, cols=2)
    widths = [CONTENT_WIDTH_DXA // 2, CONTENT_WIDTH_DXA - CONTENT_WIDTH_DXA // 2]
    set_table_geometry(meta, widths, indent_dxa=0)
    set_table_borders(meta, color=MID_GRAY, size=4)
    rows = [
        ("客户单位", "待填写"),
        ("报价单位", "待填写"),
        ("编制日期", "2026年8月19日"),
        ("文件版本", "V1.0"),
    ]
    for i, (label, value) in enumerate(rows):
        set_cell_text(meta.cell(i, 0), label, size=10, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_text(meta.cell(i, 1), value, size=10, color=DARK)
        set_cell_shading(meta.cell(i, 0), PALE_GRAY)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=22, after=8)

    add_callout(
        doc,
        "报价范围",
        "本方案仅包含软件定制开发、第三方测试、软件著作权、部署实施、培训及文档服务；不含服务器、门禁、道闸、大屏、摄像头、RFID 等硬件采购与安装。",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run("建议预算区间：120–180 万元")
    set_run_font(r, size=15, bold=True, color=GREEN)

    doc.add_page_break()


def add_summary(doc, modules: list[Module]):
    main = modules[:25]
    apps = modules[25:]
    main_days = (sum(x.days_low for x in main), sum(x.days_high for x in main))
    main_fees = (sum(x.fee_low for x in main), sum(x.fee_high for x in main))
    app1 = apps[0]
    app2 = apps[1]
    all_days = (sum(x.days_low for x in modules), sum(x.days_high for x in modules))
    software_fees = (sum(x.fee_low for x in modules), sum(x.fee_high for x in modules))
    ancillary = (12.6, 33.0)
    total = (software_fees[0] + ancillary[0], software_fees[1] + ancillary[1])

    assert main_days == (745, 965)
    assert round(main_fees[0], 1) == 89.4 and round(main_fees[1], 1) == 173.7
    assert all_days == (835, 1115)
    assert round(software_fees[0], 1) == 100.2 and round(software_fees[1], 1) == 200.7
    assert total == (112.8, 233.7)

    add_heading(doc, "一、报价总览", 1)
    add_callout(
        doc,
        "建议预算",
        "建议以 120–180 万元作为招标及商务谈判参考区间。最终合同价需在需求规格说明书（SRS）、交付边界、税务口径和工期确认后确定。",
        fill=PALE_GREEN,
        accent=GREEN,
    )

    table = doc.add_table(rows=1, cols=4)
    widths = [
        round(CONTENT_WIDTH_DXA * 0.45),
        round(CONTENT_WIDTH_DXA * 0.18),
        round(CONTENT_WIDTH_DXA * 0.18),
        CONTENT_WIDTH_DXA - round(CONTENT_WIDTH_DXA * 0.81),
    ]
    set_table_geometry(table, widths)
    set_table_borders(table)
    headers = ["费用构成", "人天区间", "金额区间（万元）", "口径说明"]
    for i, value in enumerate(headers):
        set_cell_text(table.cell(0, i), value, size=9.3, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.cell(0, i), NAVY)
    set_repeat_table_header(table.rows[0])

    rows = [
        ("主系统模块（模块1–25）", f"{main_days[0]}–{main_days[1]}", f"{main_fees[0]:.1f}–{main_fees[1]:.1f}", "按分项明细重算"),
        ("综合信息填报终端 APP", f"{app1.days_low}–{app1.days_high}", f"{app1.fee_low:.1f}–{app1.fee_high:.1f}", "含国产化适配"),
        ("移动信息终端 APP", f"{app2.days_low}–{app2.days_high}", f"{app2.fee_low:.1f}–{app2.fee_high:.1f}", "含国产化适配"),
        ("软件开发合计", f"{all_days[0]}–{all_days[1]}", f"{software_fees[0]:.1f}–{software_fees[1]:.1f}", "27个模块合计"),
        ("测试、证书、部署、培训", "—", f"{ancillary[0]:.1f}–{ancillary[1]:.1f}", "配套服务"),
        ("总估算", "—", f"{total[0]:.1f}–{total[1]:.1f}", "软件开发 + 配套服务"),
    ]
    for row_idx, values in enumerate(rows, start=1):
        cells = table.add_row().cells
        for col_idx, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.LEFT if col_idx in (0, 3) else WD_ALIGN_PARAGRAPH.CENTER
            bold = row_idx in (4, 6)
            color = NAVY if bold else DARK
            set_cell_text(cells[col_idx], value, size=9.2, bold=bold, color=color, align=align)
        if row_idx == 4:
            for cell in cells:
                set_cell_shading(cell, PALE_BLUE)
        if row_idx == 6:
            for cell in cells:
                set_cell_shading(cell, PALE_GREEN)
    set_table_geometry(table, widths)
    set_table_borders(table)

    add_callout(
        doc,
        "金额复核说明",
        "原始材料中的概览合计和“总工作量与费用汇总”与 27 个模块逐项金额不一致。本文件统一以分项明细、人天单价 1,200–1,800 元/人天重新核算，故采用总估算 112.8–233.7 万元。",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    add_heading(doc, "二、计价基础与边界", 1)
    items = [
        "人天单价按 1,200–1,800 元浮动，实际单价取决于人员级别、项目周期、交付地点及资质要求。",
        "报价覆盖需求分析、系统设计、开发编码、单元测试及分项明细中列明的配套服务。",
        "第三方测试、软件著作权、部署实施、培训与文档单独列项。",
        "国产化操作系统适配（麒麟/统信）已在两个 APP 模块中考虑相应工作量。",
        "服务器、门禁控制器、道闸、大屏、摄像头、RFID 读写器等硬件设备及安装不在本报价内。",
        "税率、发票类型、付款节点、维保期限、差旅与驻场安排尚待双方确认。",
    ]
    num_id = add_numbering_definition(doc, lvl_text="%1.", start=1, left=540, hanging=270)
    for item in items:
        add_numbered_item(doc, item, num_id, size=10.3)


def add_module_table(doc, modules: list[Module]):
    add_heading(doc, "三、分模块报价明细", 1)
    add_body(
        doc,
        "下表按 27 个软件模块列示对应技术序号、工作量和费用区间；费用按各模块估算人天与人天单价区间计算。",
        color=TEXT_GRAY,
    )

    table = doc.add_table(rows=1, cols=5)
    widths = [
        round(CONTENT_WIDTH_DXA * 0.07),
        round(CONTENT_WIDTH_DXA * 0.40),
        round(CONTENT_WIDTH_DXA * 0.16),
        round(CONTENT_WIDTH_DXA * 0.14),
        CONTENT_WIDTH_DXA - round(CONTENT_WIDTH_DXA * 0.77),
    ]
    set_table_geometry(table, widths)
    set_table_borders(table)
    headers = ["模块", "功能模块", "对应技术序号", "估算人天", "估算费用（万元）"]
    for idx, value in enumerate(headers):
        set_cell_text(table.cell(0, idx), value, size=8.8, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.cell(0, idx), NAVY)
    set_repeat_table_header(table.rows[0])

    for module in modules:
        cells = table.add_row().cells
        values = [
            str(module.number),
            module.name,
            module.reference,
            f"{module.days_low}–{module.days_high}",
            f"{module.fee_low:.1f}–{module.fee_high:.1f}",
        ]
        for idx, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.LEFT if idx == 1 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[idx], value, size=8.8, align=align)
        if module.number % 2 == 0:
            for cell in cells:
                set_cell_shading(cell, "F8FAFC")

    totals = table.add_row().cells
    total_days_low = sum(x.days_low for x in modules)
    total_days_high = sum(x.days_high for x in modules)
    total_fee_low = sum(x.fee_low for x in modules)
    total_fee_high = sum(x.fee_high for x in modules)
    values = ["", "软件开发合计", "", f"{total_days_low}–{total_days_high}", f"{total_fee_low:.1f}–{total_fee_high:.1f}"]
    for idx, value in enumerate(values):
        align = WD_ALIGN_PARAGRAPH.LEFT if idx == 1 else WD_ALIGN_PARAGRAPH.CENTER
        set_cell_text(totals[idx], value, size=9, bold=True, color=NAVY, align=align)
        set_cell_shading(totals[idx], PALE_BLUE)

    set_table_geometry(table, widths)
    set_table_borders(table)


def add_ancillary(doc):
    add_heading(doc, "四、其他必要费用", 1)
    table = doc.add_table(rows=1, cols=3)
    widths = [
        round(CONTENT_WIDTH_DXA * 0.44),
        round(CONTENT_WIDTH_DXA * 0.20),
        CONTENT_WIDTH_DXA - round(CONTENT_WIDTH_DXA * 0.64),
    ]
    set_table_geometry(table, widths)
    set_table_borders(table)
    headers = ["费用项目", "金额（万元）", "说明"]
    for idx, value in enumerate(headers):
        set_cell_text(table.cell(0, idx), value, size=9, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.cell(0, idx), NAVY)
    set_repeat_table_header(table.rows[0])

    rows = [
        ("第三方 CNAS 软件测试（GJB/Z 141-2004）", "5.0–15.0", "由具备相应资质的检测机构出具报告"),
        ("软件著作权证书办理（主系统×2 + APP×2）", "0.6–1.0", "共 4 项证书申请"),
        ("部署实施（1总部 + 34单位）", "5.0–12.0", "含安装、数据初始化、联网/离线切换及联调"),
        ("培训与文档", "2.0–5.0", "含多级用户培训、操作手册、技术文档及视频录制"),
        ("配套服务合计", "12.6–33.0", "按以上项目逐项相加"),
    ]
    for row_idx, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx == 1 else WD_ALIGN_PARAGRAPH.LEFT
            bold = row_idx == len(rows)
            set_cell_text(cells[idx], value, size=9.1, bold=bold, color=NAVY if bold else DARK, align=align)
        if row_idx == len(rows):
            for cell in cells:
                set_cell_shading(cell, PALE_BLUE)
    set_table_geometry(table, widths)
    set_table_borders(table)


def add_important_notes(doc):
    add_heading(doc, "五、重要说明", 1)
    notes = [
        "以上为纯软件定制开发及配套服务报价，不包含硬件设备采购与安装。",
        "最终报价应以正式需求规格说明书（SRS）和双方确认的功能边界为准；若功能范围、接口数量、部署单位或安全要求变化，费用与工期需相应调整。",
        "序号4“需提供同类型单位使用意见”属于投标资格项，不直接计入开发费用，但应预留商务协调时间。",
        "离线/在线自动同步涉及断点续传、数据版本、冲突检测与合并等机制，属于高风险功能，应在详细设计阶段单独评审。",
        "涉及生产厂家承诺函、软件著作权或软件产品登记证书的模块，应在投标前完成资质核验。",
        "建议在合同中明确验收标准、第三方测试范围、数据迁移边界、接口责任、源代码交付、维保期限、知识产权及保密要求。",
    ]
    num_id = add_numbering_definition(doc, lvl_text="%1.", start=1, left=540, hanging=270)
    for note in notes:
        add_numbered_item(doc, note, num_id, size=10.2)


def add_detailed_requirements(doc, modules: list[Module]):
    doc.add_page_break()
    add_heading(doc, "六、分模块详细技术要求", 1)
    add_callout(
        doc,
        "阅读说明",
        "本章保留原始技术要求内容，并将长段落拆分为便于核对的条目。每个模块前列示对应技术序号、人天和费用区间。",
        fill=PALE_BLUE,
        accent=BLUE,
    )

    for module in modules:
        h = add_heading(doc, f"模块 {module.number}｜{module.name}", 2)
        h.paragraph_format.keep_with_next = True

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.keep_with_next = True
        for label, value in (
            ("对应技术序号", module.reference),
            ("估算人天", f"{module.days_low}–{module.days_high}"),
            ("估算费用", f"{module.fee_low:.1f}–{module.fee_high:.1f} 万元"),
        ):
            r1 = p.add_run(f"{label}：")
            set_run_font(r1, size=9.3, bold=True, color=NAVY)
            r2 = p.add_run(value + "    ")
            set_run_font(r2, size=9.3, color=TEXT_GRAY)

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run("技术要求")
        set_run_font(r, size=10, bold=True, color=BLUE)

        add_complex_requirement(doc, module.requirement)

        if module.credential:
            add_callout(
                doc,
                "资质 / 承诺要求",
                module.credential,
                fill=PALE_GOLD,
                accent=GOLD,
            )


def add_confirmation(doc):
    doc.add_page_break()
    add_heading(doc, "七、商务确认", 1)
    add_body(
        doc,
        "本报价为需求阶段估算文件。双方应在正式需求、交付范围、税务口径、付款节点、工期和验收标准确认后签署最终报价单或合同。",
    )

    table = doc.add_table(rows=5, cols=2)
    widths = [CONTENT_WIDTH_DXA // 2, CONTENT_WIDTH_DXA - CONTENT_WIDTH_DXA // 2]
    set_table_geometry(table, widths)
    set_table_borders(table, color=MID_GRAY, size=5)
    rows = [
        ("报价单位（盖章）", "客户单位（签章）"),
        ("联系人：", "联系人："),
        ("联系电话：", "联系电话："),
        ("确认金额：人民币（大写）", "确认金额：¥"),
        ("日期：      年   月   日", "日期：      年   月   日"),
    ]
    for row_idx, values in enumerate(rows):
        for col_idx, value in enumerate(values):
            p = set_cell_text(
                table.cell(row_idx, col_idx),
                value,
                size=10,
                bold=row_idx == 0,
                color=NAVY if row_idx == 0 else DARK,
            )
            p.paragraph_format.space_after = Pt(14 if row_idx == 0 else 10)
            if row_idx == 0:
                set_cell_shading(table.cell(row_idx, col_idx), PALE_GRAY)
    set_table_geometry(table, widths)
    set_table_borders(table, color=MID_GRAY, size=5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("—  报价文件结束  —")
    set_run_font(r, size=9, color=TEXT_GRAY)


def build():
    source = SOURCE.read_text(encoding="utf-8")
    modules = parse_modules(source)

    doc = Document()
    add_docx_core_properties(doc)
    configure_styles(doc)
    configure_section(doc.sections[0])

    add_cover(doc)
    add_summary(doc, modules)
    add_module_table(doc, modules)
    add_ancillary(doc)
    add_important_notes(doc)
    add_detailed_requirements(doc, modules)
    add_confirmation(doc)

    # Keep tables visually separated from adjacent content.
    for table in doc.tables:
        if table._tbl.getprevious() is not None:
            pass

    doc.save(OUTPUT)
    print(f"created={OUTPUT}")
    print(f"modules={len(modules)}")
    print(f"content_width_dxa={CONTENT_WIDTH_DXA}")
    print(
        "totals="
        f"{sum(m.days_low for m in modules)}-{sum(m.days_high for m in modules)} days; "
        f"{sum(m.fee_low for m in modules):.1f}-{sum(m.fee_high for m in modules):.1f} 万元 software"
    )


if __name__ == "__main__":
    build()
